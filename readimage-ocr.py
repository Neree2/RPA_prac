import pytesseract
from docx import Document

def scanEng(imageFile):
    text=pytesseract.image_to_string(imageFile)#จะรับค่า file pic
    return text#ส่งค่าเพื่อกลับไปใช้งานต่อ

def scanThai(imageFile):
    text=pytesseract.image_to_string(imageFile, lang='tha')#จะรับค่า file pic
    return text#ส่งค่าเพื่อกลับไปใช้งานต่อ

textEng = scanEng('bookeng.jpg')
print(textEng)
print("-------------------------------------------")
textThai = scanThai('bookthai.jpg')
print(textThai)

def writeWorldFile():
    document = Document()

    document.add_heading('Scan English text from OCR')
    p= document.add_paragraph(textEng)
    p.add_run('\n\nby Neree')#ขึ้นบรรทัตใหม่

    document.add_heading('การอ่านข้อความไทยจาก OCR')
    p= document.add_paragraph(textThai)
    p.add_run('\n\nโดย เณรี')#ขึ้นบรรทัตใหม่

    document.add_page_break()#ตัวคั่นหน้า

    document.save('ocr_text.docx')

writeWorldFile()

textform= pytesseract.image_to_string('formthai.png',lang='tha')
#print(textform)

raw_text= textform.split(' ')
print(raw_text)

cleaned_text= []

for t in raw_text:
    if t != '':
        cleaned_text.append(t.strip().replace('\n',''))

print(cleaned_text)

checktext='start'
company_name=''

for s in cleaned_text:
    if 'ชื่อสถานประกอบการ' in checktext:
        company_name=s
        print(company_name)
    checktext=s